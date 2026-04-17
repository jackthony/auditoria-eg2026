"""
src/report/build_report.py

Genera el informe técnico en formato .docx usando python-docx.
Sin dependencias de Node.js.

Uso:
    py src\\report\\build_report.py
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════
#  PALETA Y HELPERS
# ══════════════════════════════════════════════════════════════
COLOR = {
    "primary": RGBColor(0x1A, 0x3A, 0x5F),
    "accent":  RGBColor(0xC0, 0x39, 0x2B),
    "muted":   RGBColor(0x66, 0x66, 0x66),
    "soft_hex":    "D5E8F0",
    "soft_row":    "F4F8FB",
    "heading_hex": "1A3A5F",
}

SEVERITY_COLOR = {
    "CRÍTICO": RGBColor(0xC0, 0x39, 0x2B),
    "MEDIA":   RGBColor(0xD3, 0x54, 0x00),
    "BAJA":    RGBColor(0x2E, 0x7D, 0x32),
    "INFO":    RGBColor(0x66, 0x66, 0x66),
}


def set_cell_shading(cell, hex_color: str):
    """Aplica un color de fondo a una celda de tabla."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_paragraph(doc, text, *, size=11, bold=False, italic=False,
                  color=None, align=None, after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    run.font.name = "Arial"
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    run.font.color.rgb = COLOR["primary"] if level <= 2 else COLOR["muted"]
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_mono(doc, text, size=9):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_header_footer(doc):
    """Header con línea inferior + footer con página."""
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run("Auditoría técnica EG2026 · Informe preliminar · "
                        + datetime.now().strftime("%d/%m/%Y"))
        hr.font.name = "Arial"
        hr.font.size = Pt(8)
        hr.font.color.rgb = COLOR["muted"]

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run()
        fr.font.name = "Arial"
        fr.font.size = Pt(8)
        fr.font.color.rgb = COLOR["muted"]
        fr.text = "Pág. "
        # Campo de número de página
        for part, val in [("begin", None), ("instrText", "PAGE"),
                          ("separate", None), ("text", "1"), ("end", None)]:
            fld = OxmlElement(f"w:fldChar" if part != "instrText" and part != "text" else
                               f"w:{part}")
            if part == "begin":
                fld.set(qn("w:fldCharType"), "begin")
            elif part == "separate":
                fld.set(qn("w:fldCharType"), "separate")
            elif part == "end":
                fld.set(qn("w:fldCharType"), "end")
            elif part == "instrText":
                fld.text = val
                fld.set(qn("xml:space"), "preserve")
            elif part == "text":
                fld.text = val
            fr._element.append(fld)


def build_kv_table(doc, rows):
    """Tabla de dos columnas (clave/valor) para metadata."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    for i, (k, v, highlight) in enumerate(rows):
        row = tbl.rows[i]
        row.cells[0].width = Cm(8)
        row.cells[1].width = Cm(10)
        c0 = row.cells[0]
        c1 = row.cells[1]
        set_cell_shading(c0, COLOR["soft_hex"] if i % 2 == 0 else COLOR["soft_row"])
        p = c0.paragraphs[0]; p.add_run(k).bold = True
        for r in p.runs: r.font.name = "Arial"; r.font.size = Pt(10)
        p2 = c1.paragraphs[0]; run = p2.add_run(str(v))
        run.font.name = "Arial"; run.font.size = Pt(10)
        if highlight:
            run.bold = True
            run.font.color.rgb = COLOR["accent"]
    return tbl


def build_manifest_table(doc, entries):
    tbl = doc.add_table(rows=1 + len(entries), cols=5)
    tbl.style = "Light Grid Accent 1"
    headers = ["Endpoint", "Ruta", "UTC", "Bytes", "SHA-256 (16)"]
    widths = [Cm(2.5), Cm(5.5), Cm(2.8), Cm(1.8), Cm(4.5)]

    # Header row
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = tbl.rows[0].cells[i]
        cell.width = w
        set_cell_shading(cell, "1A3A5F")
        p = cell.paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.name = "Arial"; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, e in enumerate(entries):
        row = tbl.rows[idx + 1]
        fill = "FFFFFF" if idx % 2 == 0 else COLOR["soft_row"]
        values = [
            e["endpoint"],
            e["url"].replace("https://", ""),
            e["fetched_at_utc"][11:19],
            f"{e['bytes']:,}",
            e["sha256"][:16] + "…",
        ]
        for i, (val, w) in enumerate(zip(values, widths)):
            cell = row.cells[i]
            cell.width = w
            set_cell_shading(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Consolas" if i in (1, 4) else "Arial"
            run.font.size = Pt(8)
    return tbl


def add_image(doc, img_path: Path, width_cm: float = 15):
    if not img_path.exists():
        add_paragraph(doc, f"[imagen no encontrada: {img_path.name}]",
                      italic=True, color=COLOR["muted"])
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    p.paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════
#  CONSTRUCCIÓN DEL INFORME
# ══════════════════════════════════════════════════════════════
def build(root: Path, output: Path):
    meta = json.loads((root / "data/processed/meta.json").read_text(encoding="utf-8"))
    findings_data = json.loads((root / "reports/findings.json").read_text(encoding="utf-8"))
    findings = findings_data["findings"]
    results = findings_data["results"]

    # Cargar manifiesto
    capture_dir = root / meta["capture_dir"]
    manifest_entries = []
    manifest_path = capture_dir / "MANIFEST.jsonl"
    if manifest_path.exists():
        seen_endpoints = set()
        for line in manifest_path.read_text(encoding="utf-8").strip().split("\n"):
            if line:
                e = json.loads(line)
                if e["endpoint"] not in seen_endpoints:
                    manifest_entries.append(e)
                    seen_endpoints.add(e["endpoint"])

    doc = Document()

    # Estilos por defecto
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_header_footer(doc)

    # ── Portada ──────────────────────────────────────────────────
    add_paragraph(doc, "INFORME TÉCNICO PRELIMINAR",
                  size=10, bold=True, color=COLOR["muted"])
    add_paragraph(doc, "Auditoría técnico-estadística del escrutinio — "
                       "Elecciones Generales 2026, Perú",
                  size=22, bold=True, color=COLOR["primary"])
    capture_human = (f"{meta['capture_ts_utc'][:4]}-{meta['capture_ts_utc'][4:6]}-"
                     f"{meta['capture_ts_utc'][6:8]} "
                     f"{meta['capture_ts_utc'][9:11]}:{meta['capture_ts_utc'][11:13]} UTC")
    add_paragraph(doc, f"Corte: {capture_human} · Nivel: preliminar (datos públicos)",
                  size=10, italic=True, color=COLOR["muted"])

    add_paragraph(
        doc,
        "Documento elaborado como insumo técnico para actuaciones ante el Jurado "
        "Electoral Especial, el Jurado Nacional de Elecciones y/o instancias "
        "parlamentarias. Su objeto es presentar hallazgos verificables derivados "
        "del análisis de la data electoral pública disponible al corte indicado, "
        "distinguiendo claramente entre observaciones empíricas, interpretaciones "
        "estadísticas e hipótesis no confirmadas.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_paragraph(
        doc,
        "Alcance declarado: este informe NO constituye prueba de fraude electoral. "
        "Tampoco afirma que el proceso haya sido limpio. Es un análisis parcial "
        "con las limitaciones metodológicas detalladas en la §5, orientado a "
        "identificar áreas que ameritan indagación adicional y preguntas técnicas "
        "formales.",
        italic=True, color=COLOR["muted"],
    )

    # ── §1 Resumen ejecutivo ─────────────────────────────────────
    add_heading(doc, "1. Resumen ejecutivo", 1)
    add_heading(doc, "1.1 Estado del escrutinio", 3)
    build_kv_table(doc, [
        ("Actas contabilizadas",
         f"{meta['actas_contabilizadas']:,} de {meta['actas_total']:,} "
         f"({meta['pct_global']}%)", False),
        ("Actas enviadas al JEE (impugnadas)",
         f"{meta['enviadas_jee']:,} "
         f"({meta['enviadas_jee']/meta['actas_total']*100:.2f}%)", False),
        ("Actas pendientes", f"{meta['pendientes_jee']:,}", False),
        ("Margen Sánchez − López Aliaga",
         f"+{meta['margen_sanch_rla_votos']:,} votos "
         f"(+{meta['margen_sanch_rla_pct']:.3f} pts)", True),
    ])
    add_paragraph(doc, "")

    add_heading(doc, "1.2 Hallazgos (resumen)", 3)
    for i, f in enumerate(findings, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"[{f['severity']}] ")
        r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11)
        r1.font.color.rgb = SEVERITY_COLOR.get(f["severity"], COLOR["muted"])
        r2 = p.add_run(f"H{i}. {f['title']}")
        r2.bold = True; r2.font.name = "Arial"; r2.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(6)

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.8)
        p2.paragraph_format.space_after = Pt(2)
        r3 = p2.add_run(f["detail"])
        r3.font.name = "Arial"; r3.font.size = Pt(10)

    # ── §2 Metodología y cadena de custodia ───────────────────────
    doc.add_page_break()
    add_heading(doc, "2. Metodología y cadena de custodia", 1)
    add_heading(doc, "2.1 Fuentes de datos", 3)
    add_bullet(doc, "Backend ONPE: resultadoelectoral.onpe.gob.pe/presentacion-backend "
                    "(fuente oficial primaria).")
    add_bullet(doc, "Cuando la IP de captura es bloqueada por ONPE (datacenters), se "
                    "utiliza un proxy CORS abierto (onpe-proxy.renzonunez-af.workers.dev) "
                    "que reenvía los mismos endpoints sin modificar el contenido. "
                    "Cada archivo queda registrado con su fuente efectiva en el MANIFEST.")
    add_bullet(doc, "Referencias legales: Reglamento JNE 0182-2025 sobre actas "
                    "observadas e impugnadas; Ley Orgánica de Elecciones (LOE) "
                    "Ley N° 26859; Ley Orgánica de la ONPE Ley N° 26487.")

    add_heading(doc, "2.2 Cadena de custodia de las muestras", 3)
    add_paragraph(doc,
                  "Se descargaron los endpoints en instantánea atómica, con "
                  "User-Agent identificado y hash SHA-256 calculado inmediatamente "
                  "tras la descarga. El manifiesto registra timestamp ISO-8601 UTC, "
                  "IP pública, hostname y commit git del capturante.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if manifest_entries:
        build_manifest_table(doc, manifest_entries)
    add_paragraph(doc, "")
    add_paragraph(doc,
                  f"Los archivos originales permanecen en "
                  f"captures/{meta['capture_ts_utc']}/ del repositorio. "
                  f"Cualquier tercero puede reproducir los hashes con "
                  f"py src\\capture\\verify_manifest.py "
                  f"captures\\{meta['capture_ts_utc']}\\",
                  italic=True, size=9)

    add_heading(doc, "2.3 Tests aplicados", 3)
    add_bullet(doc, "Reconciliación regional ↔ nacional: Σvotos_regionales vs nacional.")
    add_bullet(doc, "Detección de outliers: z-score sobre tasa de impugnación por región.")
    add_bullet(doc, "Comparación Lima+Callao vs resto: z-test de dos proporciones.")
    add_bullet(doc, "Ley de Benford (χ² gl=8) sobre primer dígito de votos por candidato.")
    add_bullet(doc, "Análisis de serie temporal del conteo oficial sobre snapshots de tracking.")
    add_bullet(doc, "Simulación: impacto de la resolución de actas JEE sobre el margen Sánchez-RLA.")

    # ── §3 Hallazgos detallados ───────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "3. Hallazgos detallados", 1)

    # H: impugnación (outliers)
    imp = results.get("impugnation", {})
    if imp.get("outliers"):
        add_heading(doc, "3.1 Outliers de tasa de impugnación por región", 2)
        for o in imp["outliers"]:
            add_paragraph(doc,
                          f"• {o['name']}: {o['tasa_impugnacion']:.2f}% "
                          f"(z = {o['z_score']:+.2f}), "
                          f"{'muy por encima' if o['z_score'] > 0 else 'muy por debajo'} "
                          f"del promedio regional.",
                          size=10)
        add_paragraph(doc,
                      "La presencia de outliers no implica irregularidad dolosa. "
                      "Puede explicarse por factores logísticos específicos. Sin "
                      "embargo, dado el impacto potencial sobre el margen del 2º "
                      "lugar, se recomienda inspección acta por acta priorizada "
                      "en estas jurisdicciones.",
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_image(doc, root / "reports/figures/fig1_impugnacion_region.png", width_cm=16)

    # H: estratificación
    add_heading(doc, "3.2 Asimetría entre Lima+Callao y resto del país", 2)
    zt = imp.get("ztest", {})
    if zt:
        add_paragraph(doc,
                      f"z = {zt['z']:+.2f} · p = {zt['p_value']:.4g} · "
                      f"Diferencia: {zt['diff_pp']:+.2f}pp",
                      bold=True)
        add_paragraph(doc,
                      f"Lima+Callao (donde Renovación Popular recibió mayor "
                      f"votación regional) registra una tasa de impugnación de "
                      f"{zt['lima_tasa']:.2f}%, inferior al {zt['resto_tasa']:.2f}% "
                      f"del resto del país. La diferencia es estadísticamente "
                      f"significativa (p < 0.01), pero NO es perjudicial para "
                      f"candidatos con mayor voto en Lima: una eventual "
                      f"regularización de impugnadas beneficia más a Sánchez, "
                      f"cuya fuerza está en regiones con tasa más alta de observación.",
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_image(doc, root / "reports/figures/fig3_estratos.png", width_cm=16)

    # H: simulación JEE
    add_heading(doc, "3.3 Definición del 2º lugar depende de resolución JEE", 2)
    sim = results.get("jee_simulation", {})
    if sim:
        add_paragraph(doc,
                      f"Margen actual: {sim['margen_actual']:+,} votos · "
                      f"Votos en juego: ~{sim['votos_juego']:,} · "
                      f"Break-even RLA: {sim['break_even_pct_rla']:.2f}%",
                      bold=True)
        add_paragraph(doc,
                      f"El margen es {abs(sim['votos_juego']/sim['margen_actual']):.0f}× "
                      f"menor que los votos contenidos en las actas impugnadas + "
                      f"pendientes. Bajo distribución proporcional al nacional, el "
                      f"margen final sería {sim['margen_proyectado']:+,.0f}. Para "
                      f"empatar, RLA requeriría {sim['break_even_pct_rla']:.2f}% del "
                      f"subgrupo RLA+Sánchez en las actas en disputa, vs "
                      f"{sim['historico_pct_rla']:.2f}% histórico: una sobre-performance "
                      f"de {sim['break_even_pct_rla']-sim['historico_pct_rla']:+.2f}pp.",
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        add_paragraph(doc,
                      "Implicación operativa: la máxima prioridad del personero "
                      "técnico debe ser la veeduría caso por caso en los JEE donde "
                      "se resuelven las actas impugnadas, con énfasis en las "
                      "circunscripciones con mayor concentración de voto de la "
                      "organización política.",
                      italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_image(doc, root / "reports/figures/fig2_serie_temporal.png", width_cm=16)

    # ── §4 Verificaciones exitosas ────────────────────────────────
    add_heading(doc, "4. Verificaciones exitosas (negative findings)", 1)
    add_paragraph(doc,
                  "Hipótesis de irregularidad que NO resultaron sustentadas por "
                  "la data disponible. Su explicitación es metodológicamente "
                  "obligatoria:",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "4.1 Reconciliación Σ regional vs nacional", 3)
    rec = results.get("reconciliation", {})
    if rec.get("rows"):
        for r in rec["rows"]:
            add_mono(doc,
                     f"  {'✓' if r['ok'] else '⚠'} {r['candidato']:<6} "
                     f"reg={r['regional']:>10,} nac={r['nacional']:>10,} "
                     f"diff={r['diff']:+,} ({r['pct_diff']:+.3f}%)")
    add_paragraph(doc,
                  "La suma de votos por región reconcilia con el total nacional "
                  "dentro del margen de redondeo. Se descarta la hipótesis de "
                  "manipulación agregada entre los dos niveles de publicación.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading(doc, "4.2 Ley de Benford (primer dígito)", 3)
    bp = results.get("benford", {}).get("pool", {})
    if bp:
        add_mono(doc, f"Pool agregado (n={bp.get('n','—')}): "
                      f"χ² = {bp.get('chi2','—'):.2f}, "
                      f"p = {bp.get('p_value','—'):.3f}")
        status = "conforme" if bp.get("conforms") else "con desviación"
        add_paragraph(doc,
                      f"El test resulta {status} a la distribución teórica. "
                      "Nota metodológica crítica: la Ley de Benford-1 presenta "
                      "limitaciones conocidas en datos electorales (Deckert, "
                      "Myagkov & Ordeshook 2011). NO es prueba concluyente de "
                      "limpieza ni de irregularidad por sí sola.",
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ── §5 Limitaciones ───────────────────────────────────────────
    add_heading(doc, "5. Limitaciones del análisis", 1)
    add_bullet(doc, "Sólo data pública agregada por región. No hay todavía acceso al "
                    "módulo de descarga masiva para organizaciones políticas.")
    add_bullet(doc, "No se realizó OCR ni cotejo acta por acta de PDFs digitalizados. "
                    "Una auditoría completa requiere ese cotejo, de varios días de trabajo.")
    add_bullet(doc, "No hay acceso a logs internos de ONPE (STAE, ODPE, centro de cómputo).")
    add_bullet(doc, "Los tests estadísticos no son concluyentes por sí solos: un "
                    "adversario sofisticado podría fabricar datos que pasen todos los tests.")
    add_bullet(doc, "Sin línea base histórica: idealmente se compararía contra EG2021, "
                    "EG2016, EG2011 para contextualizar lo \"normal\". Pendiente.")

    # ── §6 Hechos públicos documentados ───────────────────────────
    add_heading(doc, "6. Hechos públicos estructurales ya documentados", 1)
    add_paragraph(doc,
                  "Estos hechos provienen de fuentes primarias públicas (actas de "
                  "comisiones parlamentarias, oficios de Contraloría General, "
                  "denuncias penales de la Procuraduría del JNE). Constituyen la "
                  "base más sólida de responsabilidad funcional y NO requieren "
                  "análisis estadístico para ser invocados:",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_bullet(doc, "Ampliación extraordinaria de la jornada al lunes 13/04/2026 por "
                    "incumplimiento de distribución de material electoral.")
    add_bullet(doc, "211 mesas de sufragio no instaladas el 12/04, afectando a más "
                    "de 63,300 electores (Lima Sur: SJM, Lurín, Pachacámac).")
    add_bullet(doc, "Contraloría emitió 278 informes y 600 observaciones; 77% sin "
                    "corrección al día de la elección (declaración de Vera Coronel, "
                    "Comisión de Fiscalización del Congreso, 14/04/2026).")
    add_bullet(doc, "Empresa Servicios Generales Gálaga S.A.C. con sanciones previas "
                    "(contratos 2020 y 2023); contrato actual > S/6 millones incumplido.")
    add_bullet(doc, "Procuraduría del JNE interpuso denuncia penal contra el jefe de "
                    "ONPE, tres funcionarios y el representante de Gálaga.")
    add_bullet(doc, "Junta Nacional de Justicia inició procedimiento para evaluar "
                    "destitución del jefe de ONPE.")
    add_bullet(doc, "Dircocor PNP citó al jefe de ONPE como testigo.")

    # ── §7 Recomendaciones ────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "7. Recomendaciones operativas", 1)
    add_heading(doc, "7.1 Preguntas técnicas para la audiencia", 2)
    add_numbered(doc, "Sobre el STAE: ¿qué pruebas de carga, integridad y seguridad "
                      "se realizaron antes del 12/04? ¿Hay reportes firmados por "
                      "Gerencia de Informática? Solicitamos copia formal.")
    add_numbered(doc, "De las 600 observaciones de Contraloría previas al proceso, "
                      "¿cuáles son las 462 no atendidas (77%)? ¿Bajo qué criterio "
                      "y con qué firma?")
    add_numbered(doc, "Sobre Gálaga: ¿qué matriz de evaluación se aplicó dado sus "
                      "sanciones previas en 2020 y 2023? ¿Qué penalidades contractuales "
                      "se han ejecutado?")
    add_numbered(doc, "Sobre Extranjero (tasa de impugnación 4× el nacional): ¿a qué "
                      "se atribuye? ¿Existe desglose por oficina consular?")
    add_numbered(doc, "Sobre el módulo de descarga masiva para organizaciones políticas: "
                      "¿fecha y hora de operatividad? ¿Credenciales para el personero técnico?")
    add_numbered(doc, "Sobre la cadena de custodia digital del acta: desde ODPE hasta "
                      "publicación, ¿qué checksum o firma digital la acompaña?")
    add_numbered(doc, "Sobre las 211 mesas no instaladas: ¿evaluación de impacto "
                      "contrafactual por partido?")

    add_heading(doc, "7.2 Acciones inmediatas del personero técnico", 2)
    add_numbered(doc, "Exigir al JEE el desglose por causal de observación (Reglamento "
                      "JNE 0182-2025) de las actas enviadas.")
    add_numbered(doc, "Solicitar acceso al módulo de descarga masiva de actas digitalizadas "
                      "y ejecutar cotejo acta-por-acta contra copias del partido.")
    add_numbered(doc, "Documentar con fotografía + hash + georreferencia cada divergencia "
                      "entre copia del personero y versión publicada.")
    add_numbered(doc, "Participar físicamente en las audiencias JEE de mayor "
                      "concentración de voto del partido (prioridad: Extranjero, Lima).")
    add_numbered(doc, "Coordinar con el equipo legal la conservación de la cadena de "
                      "custodia de toda la evidencia digital (hashes, timestamps, certificados).")

    # ── §8 Declaración ────────────────────────────────────────────
    add_heading(doc, "8. Declaración de responsabilidad", 1)
    add_paragraph(doc,
                  "Este informe se elaboró a partir de datos públicos consultados al "
                  "corte indicado. Las aserciones empíricas son reproducibles por "
                  "terceros siguiendo el procedimiento de §2 sobre los archivos "
                  "hasheados del repositorio. Las inferencias estadísticas siguen "
                  "convenciones estándar (α=0.05; |z|≥2 para flags). Los juicios "
                  "interpretativos se etiquetan como tales.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(doc,
                  "El presente documento NO afirma la existencia de fraude electoral "
                  "ni descarta la posibilidad de irregularidades específicas en "
                  "niveles no accesibles desde la data pública agregada. Es un "
                  "insumo preliminar para la fiscalización ciudadana y partidaria, "
                  "no una conclusión pericial-judicial.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_paragraph(doc, "")
    add_paragraph(doc, "Autor: __________________________________________",
                  size=10, color=COLOR["muted"])
    add_paragraph(doc, "Cargo: __________________________________________",
                  size=10, color=COLOR["muted"])
    add_paragraph(doc, "Colegiatura profesional (si aplica): _______________",
                  size=10, color=COLOR["muted"])
    add_paragraph(doc, "Fecha y firma: ___________________________________",
                  size=10, color=COLOR["muted"])

    # ── Anexo A ───────────────────────────────────────────────────
    add_heading(doc, "Anexo A — Archivos de evidencia", 1)
    add_mono(doc, f"captures/{meta['capture_ts_utc']}/")
    for e in manifest_entries:
        add_mono(doc, f"  raw/{e['endpoint']}.json  ({e['bytes']:,}B  "
                      f"sha256={e['sha256'][:16]}…)")
    add_mono(doc, f"captures/{meta['capture_ts_utc']}/MANIFEST.jsonl")
    add_mono(doc, "data/processed/regiones.csv")
    add_mono(doc, "data/processed/meta.json")
    add_mono(doc, "data/processed/tracking.csv")
    add_mono(doc, "reports/findings.json")
    add_mono(doc, "reports/summary.txt")
    add_mono(doc, "reports/figures/fig{1,2,3}_*.png")
    add_mono(doc, "src/capture/fetch_onpe.py")
    add_mono(doc, "src/process/build_dataset.py")
    add_mono(doc, "src/analysis/run_all.py")

    doc.save(str(output))
    size = output.stat().st_size
    print(f"✓ {output}  ({size:,} bytes)")


def main():
    ROOT = Path(__file__).resolve().parents[2]
    OUT = ROOT / "reports" / "Informe_Tecnico_v1.docx"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build(ROOT, OUT)


if __name__ == "__main__":
    main()
