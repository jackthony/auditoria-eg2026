from docx import Document
from pathlib import Path

src = Path('reports/Informe_Tecnico_v1.docx')
dst = Path('reports/Informe_Tecnico_v1.txt')

d = Document(src)
out = []
for p in d.paragraphs:
    if p.text.strip():
        out.append(p.text)
for i, t in enumerate(d.tables):
    out.append(f'\n--- TABLA {i+1} ---')
    for row in t.rows:
        out.append(' | '.join(c.text for c in row.cells))

dst.write_text('\n'.join(out), encoding='utf-8')
print(f'OK: {dst} ({dst.stat().st_size} bytes)')
